from bs4 import BeautifulSoup, NavigableString
import requests
import docx
import sys
import argparse
import os
import html
import re


class Comment:
    # levels store indentation info for children comments' indentation
    # siteTable is header post's parent-id - it's -2 so header's would be -1
    levels = {"siteTable": -2}

    def __init__(self, data, parent=None):
        self.id = data['data-fullname']
        # comments received via POST calls don't have parent ids in their
        # contents, but provide them separately
        if not parent:
            # the actual link present is in the form of siteTable_PARENT-ID
            if len(data.parent['id']) > 10:
                self.parent = data.parent['id'][10:]
            # ...with the exception of the header post, which is siteTable
            else:
                self.parent = data.parent['id']
        else:
            self.parent = parent
        # messages from deleted profiles don't have this attribute
        try:
            self.author = '[{}] '.format(data['data-author'])
        # so this is a placeholder author for those messages
        except KeyError:
            self.author = '[DELETED] '
        try:
            # all paragraphs of a single message
            self.text = [a.contents for a in [a for a in data.find(
                    'div', class_='md').contents if a != '\n']]
        except AttributeError:
            self.text = '<empty message>'
        try:
            self.level = Comment.levels[self.parent] + 1
        # deleted messages won't be present in levels, deleted() adds it
        except KeyError:
            deleted(self.parent)
            self.level = Comment.levels[self.parent] + 1
        # store level information for further use by message's children
        Comment.levels[self.id] = self.level

    def convert(self, document):
        # reduce level_size manually if there're too many levels
        level_size = docx.shared.Mm(5)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = \
            docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        # header should appear at the same level as 0th level message
        paragraph.paragraph_format.left_indent = level_size * self.level if \
            self.level >= 0 else 0
        paragraph.add_run(self.author).bold = True
        if isinstance(self.text, list):
            # iterate over all message's paragraphs
            for a in range(len(self.text)):
                # if a message contains several paragraphs, every paragraph
                # after the first should be separated by a line break
                if a > 0:
                    paragraph.add_run('\n')
                # recursive function to grab all formatting tags for a message
                for b in self.text[a]:
                    self.__stick(b, paragraph)
        else:
            self.__stick(self.text, paragraph)
        return document

    # used on the bottom level of stick function
    # applies all grabbed formatting options
    def __decorate(self, tag, paragraph, decor):
        text = str(tag).replace('\n', '')
        if decor['link']:
            add_hyperlink(text, paragraph, decor['link'])
        else:
            tree = paragraph.add_run(text)
            if decor['italics']:
                tree.italic = True
            if decor['strong']:
                tree.bold = True
            if decor['super']:
                tree.font.superscript = True
            if decor['strike']:
                tree.font.strike = True

    def __stick(self, tag, paragraph, schema={'link': 0, 'color': 0,
                'italics': 0, 'strong': 0, 'super': 0, 'strike': 0}):
        # NavigableString = tag has no children = bottom level = use decorate
        if isinstance(tag, NavigableString) or isinstance(tag, str):
            self.__decorate(tag, paragraph, schema)
        else:
            str_tag = str(tag)[:3]
            decor = schema.copy()
            # the tags supported
            if str_tag == '<br':
                paragraph.add_run('\n')
            elif str_tag == '<a ':
                decor['link'] = tag['href']
            elif str_tag == '<em':
                decor['italics'] = True
            elif str_tag == '<st':
                decor['strong'] = True
            elif str_tag == '<li':
                paragraph.add_run('\n - ')
            elif str_tag == '<p>':
                if str(tag.parent)[:3] == '<bl':
                    paragraph.add_run('\n > ')
            elif str_tag == '<su':
                decor['super'] = True
            elif str_tag == '<de':
                decor['strike'] = True
            for a in tag.contents:
                self.__stick(a, paragraph, decor)


# returns true if a tag is a comment or an object of interest
def checker(tag):
    return tag.has_attr('data-type') and tag['data-type'] in [
        'comment', 'morerecursion', 'morechildren']


# resolves issues with deleted comments by finding their level
def deleted(deleted_id):
    deleted_info = requests.get('https://www.reddit.com/api/info.json?id={}'.
                                format(deleted_id), headers=agent)
    if deleted_info.status_code != requests.codes.ok:
        print('ERROR: {} while trying to retrieve {}'.format(
            deleted_info.status_code, results.input))
        sys.exit()
    Comment.levels[deleted_id] = deleted_info.json()['data']['dist']


# internet solution, as docx does not have pre-built hyperlink method
def add_hyperlink(text, paragraph, link):
    part = paragraph.part
    r_id = part.relate_to(link, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# processes data and stitches together initially loaded messages,
# 'load more comments' from POST calls and 'continue this thread' links
def autoglue(parent_id, data, document, is_post=False):
    raw_data = data.find_all(checker)
    # 1st message of every new file is a repetition and thus must be ignored
    # in 'continue this thread' links
    if is_post:
        comments = [Comment(comment, parent_id) if comment['data-type'] ==
                    'comment' else comment for comment in raw_data if
                    comment['data-fullname'] != parent_id]
    else:
        comments = [Comment(comment) if comment['data-type'] == 'comment' else
                    comment for comment in raw_data if
                    comment['data-fullname'] != parent_id]
    for comment in comments:
        # use conversion on messages and more autoglue on further file breaks
        if isinstance(comment, Comment):
            document = comment.convert(document)
        else:
            if comment['data-type'] == 'morerecursion':
                deepthread = requests.get(
                    'https://www.reddit.com/' +
                    comment.find('a')['data-href-url'], headers=agent)
                if deepthread.status_code != requests.codes.ok:
                    print('ERROR: {} while trying to retrieve {}'.format(
                            deepthread.status_code, results.input))
                    sys.exit()
                soup_deepthread = BeautifulSoup(deepthread.text, "lxml")
                trgt = soup_deepthread.find(attrs={'data-type': 'comment'})
                if trgt is not None:
                    document = autoglue(
                        trgt['data-fullname'], soup_deepthread, document)
            # if it's not 'continue this thread', it's 'load more messages'
            else:
                button = comment.find("a", class_="button")
                button = button['onclick'][
                    button['onclick'].find("(")+1:button['onclick'].find(")")]\
                    .replace("'", "").split(', ')
                button_children = button[3]
                button_depth = button[4]
                button_id = 't1_' + button_children.split(',')[0]
                button_link_id = button[1]
                post = requests.post(
                    'https://www.reddit.com/api/morechildren', headers=agent,
                    data={
                        'children': button_children, 'depth': button_depth,
                        'id': button_id, 'limit_children': False, 'link_id':
                        button_link_id, 'r': subreddit, 'renderstyle': 'html',
                        'sort': 'confidence'})
                posted = [
                    data['data'] for data in post.json()["jquery"][10][3][0]]
                # add every comment received to the .docx document
                for item in posted:
                    document = autoglue(item['parent'], BeautifulSoup(
                        html.unescape(item['content']), "lxml"), document,
                        True)
    return document


parser = argparse.ArgumentParser(description="converts a reddit thread to \
    .docx")
parser.add_argument('-i', '--input',
                    help='link to a reddit thread, ex. \
                    "https://www.reddit.com/r/SUBREDDIT/comments/THREAD_ID"',
                    required='True')
parser.add_argument('-o', '--output',
                    help='destination and name for the .docx, ex. \
                        "D:/foo.docx"')
results = parser.parse_args(sys.argv[1:])

agent = {'User-agent': 'Reddit_to_docx'}

# page setup
doc = docx.Document()
doc.sections[0].orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
doc.sections[0].page_height, doc.sections[0].page_width = \
    doc.sections[0].page_width, doc.sections[0].page_height
margin = docx.shared.Mm(5)
doc.sections[0].left_margin = margin
doc.sections[0].right_margin = margin
doc.sections[0].top_margin = margin
doc.sections[0].bottom_margin = margin

first = requests.get(results.input, headers=agent)
if first.status_code != requests.codes.ok:
    print('ERROR: {} while trying to retrieve {}'.format(
        first.status_code, results.input))
    sys.exit()
file = BeautifulSoup(first.text, "lxml")
# get subreddit, will be needed for 'load more comments' calls
subreddit = file.find(attrs={'data-subreddit': True})['data-subreddit']
# parse header
title = file.find("p", class_="title").a.string
head = doc.add_paragraph()
head.style = doc.styles['Title']
header = file.find(attrs={'data-type': 'link'})
add_hyperlink(title, head, header['data-url'])
header = Comment(header)
doc = header.convert(doc)
# parse messages
doc = autoglue("siteTable", file, doc)

if results.output is None:
    doc.save(os.path.join(os.path.dirname(os.path.realpath(
        sys.argv[0])), re.sub('[^A-z0-9 -]', '', title)+'.docx'))
else:
    doc.save(results.output)
